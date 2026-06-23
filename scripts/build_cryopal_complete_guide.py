from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WALKTHROUGH_MEDIA = Path("/tmp/cryopal_walkthrough_extract/word/media")
OUTPUT_DOCX = ROOT / "docs" / "CryoPal_tomo_Complete_Guide.docx"
LOGO_PATH = ROOT / "cryoet_organizer" / "assets" / "CryoPal_tomo_logo.png"


HEADING_BLUE = RGBColor(0x2E, 0x74, 0xB5)
HEADING_DARK = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)


@dataclass(frozen=True)
class WalkthroughSection:
    title: str
    goal: str
    steps: tuple[str, ...]
    checks: tuple[str, ...]
    image_count: int
    captions: tuple[str, ...]
    tips: tuple[str, ...] = ()


@dataclass(frozen=True)
class ConceptSection:
    title: str
    paragraphs: tuple[str, ...]
    bullets: tuple[str, ...] = ()


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_run_font(run, *, name: str = "Calibri", size: int = 11, bold: bool = False, italic: bool = False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def style_document(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Heading 1", 16, HEADING_BLUE, 16, 8),
        ("Heading 2", 13, HEADING_BLUE, 12, 6),
        ("Heading 3", 12, HEADING_DARK, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if not header_para.text:
            run = header_para.add_run("CryoPal_tomo Complete Guide")
            set_run_font(run, size=9, color=MUTED)

        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if not footer_para.text:
            prefix = footer_para.add_run("Page ")
            set_run_font(prefix, size=9, color=MUTED)
            add_page_number(footer_para)


def add_cover(document: Document) -> None:
    if LOGO_PATH.exists():
        para = document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(str(LOGO_PATH), width=Inches(2.2))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("CryoPal_tomo")
    set_run_font(run, size=24, bold=True, color=HEADING_DARK)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Conceptual Guide and Hands-on Walkthrough for Cryo-ET Processing")
    set_run_font(run, size=15, color=HEADING_BLUE)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        "Comprehensive user documentation",
        "Organized as abstract, conceptual reference, and screenshot-based walkthrough",
        "Prepared on 22 June 2026",
    ):
        run = meta.add_run(line + "\n")
        set_run_font(run, size=11, color=MUTED)

    document.add_page_break()


def add_prefatory_section(document: Document) -> None:
    document.add_heading("Abstract", level=1)
    abstract = (
        "CryoPal_tomo is a project-centered desktop application for organizing, launching, tracking, "
        "and documenting cryo-electron tomography processing workflows. It acts as a coordination "
        "layer across datasets, file roles, processing jobs, tomogram inspection, particle utilities, "
        "and reproducible settings. This document combines a conceptual description of CryoPal_tomo with a "
        "practical screenshot-based walkthrough so that new users can both understand the software and "
        "apply it directly to real cryo-ET projects."
    )
    document.add_paragraph(abstract)

    document.add_heading("Preface", level=1)
    preface_paragraphs = (
        "This guide is intended for users who want to process cryo-ET datasets in a structured, "
        "traceable, and reusable way. It focuses on how CryoPal_tomo supports real project work rather than "
        "on internal implementation details.",
        "The document is organized in four parts. After the opening pages and the table of contents, "
        "the reader first gets a compact Quick Start Guide for immediate use. This is followed by a longer "
        "conceptual section that explains CryoPal_tomo's main ideas, tabs, and workflows. The final part is a "
        "hands-on walkthrough that follows a typical processing journey with screenshots and practical, "
        "stepwise instructions.",
    )
    for paragraph in preface_paragraphs:
        document.add_paragraph(paragraph)


def add_part_heading(document: Document, title: str, subtitle: str | None = None) -> None:
    document.add_page_break()
    para = document.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title)
    set_run_font(run, size=18, bold=True, color=HEADING_DARK)
    if subtitle:
        sub = document.add_paragraph()
        sub.paragraph_format.space_after = Pt(10)
        run = sub.add_run(subtitle)
        set_run_font(run, size=11, italic=True, color=MUTED)


def add_table_of_contents(document: Document) -> None:
    document.add_page_break()
    document.add_heading("Contents", level=1)
    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(8)
    run = note.add_run(
        "If the table of contents does not populate automatically in Word, click inside it and update the field."
    )
    set_run_font(run, size=10, italic=True, color=MUTED)

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Table of contents"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def add_quick_start(document: Document, images: list[Path]) -> None:
    add_part_heading(
        document,
        "Part I. Quick Start Guide",
        "The shortest practical route from a new CryoPal_tomo project to processed and curated tomograms.",
    )

    document.add_paragraph(
        "This section is intentionally minimal. If you only want to get started quickly, follow these steps first "
        "and return to the longer conceptual or walkthrough sections afterwards when you need more detail."
    )

    document.add_heading("Quick Start in seven steps", level=1)
    quick_steps = (
        "Open CryoPal_tomo and create a new project through File > New Project. Save the resulting .cryopal.json file immediately.",
        "Go to Project Overview and choose Dataset actions > Add dataset for processing. Fill in the dataset name, raw frames folder, MDOC folder, processing folder, and the key imaging parameters.",
        "Move to Processing: WARP, select the dataset, and start with the WarpTools preprocessing jobs you need. Use the command preview to verify that the resolved paths and main parameters look correct before running anything.",
        "If you prefer cluster execution, define a Slurm profile in Settings first and then submit the WARP jobs from CryoPal_tomo instead of running them locally.",
        "After tomograms have been reconstructed, open Tomogram Gallery. Review the thumbnails, rate the tomograms, and assign tags so that good or interesting tilt series can be found again quickly.",
        "Use Multi selection in the gallery to collect the best tomograms and send them into the TS processing list for downstream jobs such as segmentation, denoising, or template matching.",
        "Use Settings > Check paths and the export functions under File whenever you want to validate the project state or create machine-readable documentation of paths and job history.",
    )
    add_manual_numbered_steps(document, quick_steps)

    add_figure(document, images[0], "Figure 1. Main CryoPal_tomo window at startup.")
    add_figure(document, images[2], "Figure 2. WARP processing as the main preprocessing entry point.")
    add_figure(document, images[10], "Figure 3. Tomogram Gallery for curation, annotation, and TS selection.")


def add_manual_numbered_steps(document: Document, steps: tuple[str, ...]) -> None:
    for index, step in enumerate(steps, start=1):
        para = document.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.18)
        para.paragraph_format.first_line_indent = Inches(-0.18)
        para.paragraph_format.space_after = Pt(4)
        number_run = para.add_run(f"{index}. ")
        set_run_font(number_run, bold=True)
        text_run = para.add_run(step)
        set_run_font(text_run)


def add_bullet_points(document: Document, bullets: tuple[str, ...]) -> None:
    for bullet in bullets:
        document.add_paragraph(bullet, style="List Bullet")


def conceptual_sections() -> tuple[ConceptSection, ...]:
    return (
        ConceptSection(
            title="1. What CryoPal_tomo is designed to do",
            paragraphs=(
                "CryoPal_tomo is a project-centered desktop application for organizing and processing cryo-electron "
                "tomography datasets. Its main purpose is not to replace WarpTools, MTools, MCore, PyTom, MemBrain, "
                "or other external software, but to coordinate how those tools are used in a real project.",
                "In practical terms, CryoPal_tomo helps users keep track of datasets, file paths, settings files, job "
                "parameters, job history, tomogram curation, particle-analysis products, and reusable settings. This "
                "becomes especially valuable when several datasets, many tilt series, or multiple downstream tools are "
                "involved.",
            ),
            bullets=(
                "Project-based dataset organization",
                "Central path and file-role resolution",
                "Built-in execution support for local runs and Slurm submission",
                "Tomogram curation through thumbnails, tags, and ratings",
                "Reusable defaults, environments, viewer rules, and export functions",
            ),
        ),
        ConceptSection(
            title="2. How CryoPal_tomo organizes a project",
            paragraphs=(
                "A CryoPal_tomo project is stored as a .cryopal.json file and acts as the container for the full working "
                "state of one cryo-ET study. Each project can include several datasets, and each dataset stores the "
                "paths and metadata needed to make later processing reproducible.",
                "A dataset typically includes the raw frames folder, MDOC folder, optional gain file, processing "
                "folder, pixel size, exposure, tomogram dimensions, and later also references to settings files, "
                "tomograms, thumbnails, and job history entries. Because this information is stored centrally, the "
                "same dataset can be reused consistently in WARP, TS jobs, particle jobs, and gallery curation.",
            ),
        ),
        ConceptSection(
            title="3. Project Overview and dataset setup",
            paragraphs=(
                "Project Overview is where most users begin. It is the place to add new datasets, import already "
                "processed datasets, inspect the dataset table, and remove datasets from the project when needed. "
                "A good setup here saves a great deal of time later because the rest of CryoPal_tomo depends on the paths "
                "and metadata entered at this stage.",
                "One of the most important options during dataset creation is Unify mdoc names. When enabled, CryoPal_tomo "
                "copies and renames the MDOC files into a standardized naming scheme so that each tilt series receives "
                "a predictable and unique identifier. This is useful when original MDOC names are inconsistent or "
                "difficult to match reliably. If your source MDOC names are already stable and you want to preserve "
                "them exactly, this option should remain disabled.",
                "The other MDOC-related checkboxes are equally practical. Ignore override.mdoc prevents CryoPal_tomo from "
                "treating override files as primary tilt-series records, while Ignore custom.mdoc helps exclude "
                "alternative or auxiliary MDOC-like files that should not be interpreted as true tilt-series inputs.",
            ),
            bullets=(
                "Use Add dataset for processing for fresh projects starting from raw data.",
                "Use Import already processed dataset when Warp settings and outputs already exist.",
                "Double-clicking a dataset entry opens details and is one of the quickest sanity checks in the software.",
            ),
        ),
        ConceptSection(
            title="4. File Registry, Tomogram Gallery, and the TS processing list",
            paragraphs=(
                "CryoPal_tomo relies heavily on the File Registry. The File Registry defines how important file roles are "
                "resolved, for example tomograms, aligned stacks, angle files, tomostars, MDOCs, and user-defined "
                "roles such as segmentations or masks. This means that CryoPal_tomo does not need a separate ad-hoc search "
                "logic in every tab; instead, several tabs can consult the same resolution rules.",
                "The Tomogram Gallery sits on top of this path knowledge and adds curation. It gives users a thumbnail "
                "overview of reconstructed tomograms, together with ratings, tags, dataset context, and links to "
                "associated files. This turns the gallery into a QC and selection layer rather than a passive image browser.",
                "The TS processing list is one of the most important bridge concepts in CryoPal_tomo. It is the working "
                "selection of tilt series that downstream TS jobs and some custom jobs operate on. In practice, users "
                "often review tomograms in the gallery, select the best or most relevant tilt series, and then send "
                "that curated subset into the TS processing list for segmentation, denoising, template matching, or "
                "other per-TS tasks.",
            ),
            bullets=(
                "Use the File Registry to teach CryoPal_tomo about new file roles.",
                "Use the Gallery to decide which tilt series are worth continued work.",
                "Use the TS processing list as the handoff from curation to downstream TS-based processing.",
            ),
        ),
        ConceptSection(
            title="5. The processing tabs and what each one is for",
            paragraphs=(
                "CryoPal_tomo separates processing work by job family. Processing: WARP is mainly for WarpTools jobs and "
                "therefore covers the path from raw data toward reconstructed tomograms. Processing: M is for MTools "
                "and MCore workflows based on M populations. Processing: TS jobs contains operations that act on "
                "selected tilt series rather than on full datasets or particle STAR files. Processing: Particle jobs "
                "is centered on particle STAR files and related utilities, and Processing: Custom jobs exists for "
                "lab-specific tools that still benefit from CryoPal_tomo's metadata and file-resolution logic.",
                "The Shortcuts tab is intentionally different from the processing tabs. It is meant for convenience "
                "actions such as launching a GUI, moving into a working directory, or starting a short command "
                "sequence, without treating those actions as formal project processing history.",
            ),
        ),
        ConceptSection(
            title="6. Running, scheduling, and submitting jobs",
            paragraphs=(
                "Most CryoPal_tomo processing tabs follow the same basic interaction model. The user selects a job, "
                "reviews the generated command preview, adjusts parameters, and then decides whether to copy the "
                "command, run it locally, submit it to Slurm, or schedule it for later.",
                "Local execution can use named environments that are managed centrally in Settings. Slurm execution "
                "uses named Slurm profiles, which can include site-specific flags and setup commands. Scheduled jobs "
                "allow users to prepare larger batches deliberately and then run or submit them later in a controlled order.",
                "Job history is not just archival. It is also part of the workflow: it allows users to inspect what "
                "happened, see current queue state, reopen details, and in many cases reuse earlier parameters.",
            ),
            bullets=(
                "Always trust the command preview more than the visual form alone.",
                "Use saved defaults for recurring parameter values.",
                "Use scheduled runs when you want to build a queue before launching work.",
            ),
        ),
        ConceptSection(
            title="7. Particle utilities, exports, and project maintenance",
            paragraphs=(
                "Particle jobs extend CryoPal_tomo beyond command launching into lightweight analysis and STAR-file "
                "utilities. Export particles, merge or split STAR files, intersect particle sets, perform distance "
                "cleaning, and generate abundance or classification-convergence plots can all be handled in the "
                "particle tab.",
                "CryoPal_tomo also provides maintenance tools that become increasingly important in larger projects. "
                "Check paths validates which files and directories CryoPal_tomo can currently resolve. Export job history "
                "and Export file paths produce machine-readable CSV outputs for documentation or downstream tools. "
                "Settings bundles allow reusable defaults, environments, custom jobs, and viewer rules to be shared "
                "across projects without replacing the actual project file.",
            ),
        ),
        ConceptSection(
            title="8. What a new user should keep in mind",
            paragraphs=(
                "CryoPal_tomo is easiest to learn if it is treated as a structured workflow companion rather than as a "
                "single giant form. Start with one dataset, make sure the paths are right, run a small WARP step, "
                "inspect the results in the gallery, and only then expand into TS jobs, particle jobs, or custom roles.",
                "Once the dataset setup, file registry, and main defaults are correct, the rest of CryoPal_tomo becomes "
                "considerably easier to use. Most later convenience comes from that early investment in clean project structure.",
            ),
            bullets=(
                "Do not rush dataset setup; wrong paths at the beginning create confusion everywhere else later.",
                "Use the command preview as a teaching tool: it shows how CryoPal_tomo translates GUI choices into real commands.",
                "Treat tags, ratings, and history not as optional extras, but as part of your reproducible workflow.",
            ),
        ),
    )


def add_conceptual_guide(document: Document) -> None:
    add_part_heading(
        document,
        "Part II. Conceptual Guide",
        "A condensed explanation of CryoPal_tomo's main ideas, interface structure, and workflow logic.",
    )
    for section in conceptual_sections():
        document.add_heading(section.title, level=1)
        for paragraph in section.paragraphs:
            document.add_paragraph(paragraph)
        if section.bullets:
            add_bullet_points(document, section.bullets)


def natural_image_list() -> list[Path]:
    def key(path: Path) -> int:
        match = re.search(r"(\d+)", path.stem)
        return int(match.group(1)) if match else 0

    return sorted(WALKTHROUGH_MEDIA.glob("image*.png"), key=key)


def add_figure(document: Document, image_path: Path, caption: str) -> None:
    fig_para = document.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.add_run().add_picture(str(image_path), width=Inches(5.9))

    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    run = cap.add_run(caption)
    set_run_font(run, size=10, italic=True, color=MUTED)


def walkthrough_sections() -> tuple[WalkthroughSection, ...]:
    return (
        WalkthroughSection(
            title="1. Interface overview",
            goal="Use the first screen to orient yourself before creating or loading data.",
            steps=(
                "Start CryoPal_tomo and take a moment to identify the three main interface areas: the top menu, the sidebar, and the main content area.",
                "Use the sidebar to switch between the major work modes. Project Overview is for datasets, Tomogram Gallery is for curation, and the processing tabs are grouped by job type.",
                "Open the File menu and confirm where project creation, opening, saving, and CSV exports are located.",
                "Open the Settings menu and confirm where environments, Slurm profiles, viewer defaults, and default parameters are managed.",
            ),
            checks=(
                "Before moving on, make sure you know where Project Overview, Processing: WARP, Tomogram Gallery, and Settings are located.",
            ),
            image_count=2,
            captions=(
                "Figure 1. CryoPal_tomo main window after startup.",
                "Figure 2. Top-bar actions for file handling and global settings.",
            ),
            tips=(
                "A new user should spend a few minutes simply clicking through the sidebar before trying to process data. The interface becomes much easier once you recognize which tab is responsible for which kind of task.",
            ),
        ),
        WalkthroughSection(
            title="2. Creating a project and adding the first dataset",
            goal="Create the project container and register the raw input data so that CryoPal_tomo can resolve everything else from there.",
            steps=(
                "Choose File > New Project and save the new .cryopal.json file right away.",
                "Open Project Overview and choose Dataset actions > Add dataset for processing.",
                "Fill in the dataset name, sample, comment, raw frames folder, MDOC folder, optional gain file, processing folder, pixel size, exposure, and tomogram dimensions.",
                "Review the MDOC-related checkboxes. Enable or disable Unify mdoc names, Ignore override.mdoc, and Ignore custom.mdoc according to your acquisition setup.",
                "Click Add dataset and confirm that the dataset appears in the overview table with the expected number of tilt series.",
                "If your Warp project already exists, use Dataset actions > Import already processed dataset instead and provide the frame-series and tilt-series settings files.",
            ),
            checks=(
                "Double-click the new dataset entry and verify that the stored paths and metadata look correct before starting processing.",
            ),
            image_count=0,
            captions=(),
            tips=(
                "If you are uncertain whether Unify mdoc names should be enabled, ask yourself whether the original MDOC filenames are already consistent and unique. If they are messy or inconsistent, standardizing them usually makes later matching more robust.",
                "The processing folder should be a place where CryoPal_tomo is allowed to create and organize dataset-specific outputs over time. Think of it as your working area, not just as a temporary export target.",
            ),
        ),
        WalkthroughSection(
            title="3. Preprocessing data with WarpTools",
            goal="Run the WarpTools preprocessing steps from inside CryoPal_tomo instead of building each command manually.",
            steps=(
                "Open Processing: WARP and select the dataset you want to process.",
                "Choose whether the job belongs to the frame-series or tilt-series branch, then select the exact WarpTools job type.",
                "Review the generated command preview before running anything. This is the fastest way to catch wrong paths or inherited defaults that do not belong to the current run.",
                "Fill in or adjust the parameters you need. If you notice that a parameter should always have the same value in your lab, stop here and save it later through Settings > Set default parameters.",
                "Choose the execution mode. For quick local runs, use Run locally and select an environment if necessary. For cluster runs, choose Submit to Slurm and select a Slurm profile.",
                "If you are not ready to run yet, click Schedule command instead. Scheduled entries remain in the job history until you trigger them.",
                "When you want to process several scheduled jobs together, use Run scheduled jobs or Submit scheduled jobs to Slurm from the history section.",
            ),
            checks=(
                "Blue history entries indicate completed jobs, green indicates running or queued-to-run activity, and grey indicates jobs that are scheduled but not yet started.",
                "When in doubt, trust the command preview before you trust the form: it shows what CryoPal_tomo is really about to execute.",
            ),
            image_count=8,
            captions=(
                "Figure 3. WARP processing tab with dataset and job selection.",
                "Figure 4. Default-parameter editor for WARP job types.",
                "Figure 5. Environment-management dialog used for local execution.",
                "Figure 6. WARP command preview with local and Slurm execution choices.",
                "Figure 7. Scheduled WARP jobs in the history view.",
                "Figure 8. Collective Slurm-submission dialog for a scheduled batch.",
                "Figure 9. Queue handling for additional scheduled jobs.",
                "Figure 10. Color-coded WARP job history during execution.",
            ),
            tips=(
                "For a first dataset, do not schedule a very large batch immediately. Run one small representative job first and inspect whether the outputs land where you expect.",
                "Default parameters are especially useful for lab-standard Slurm settings, frequently reused pixel-size related values, and preferred local environments.",
            ),
        ),
        WalkthroughSection(
            title="4. Annotating and organizing tomograms in the gallery",
            goal="Turn reconstructed tomograms into a curated working set for downstream analysis.",
            steps=(
                "Open Tomogram Gallery after reconstruction jobs have produced thumbnails and associated tomograms.",
                "Click a thumbnail to inspect its metadata in the right-hand sidebar, then open the associated .mrc file if you want to review the volume externally.",
                "Assign a rating and add tags that describe why a tomogram is useful, problematic, or interesting.",
                "Use the filter bar at the top to search by dataset, minimum rating, include-tags, and exclude-tags.",
                "Activate Multi selection if you want to mark several good tomograms at once.",
                "Once your filtered set looks right, send the selected tilt series into the TS processing list for downstream jobs.",
            ),
            checks=(
                "Use tags to encode decisions you will want to revisit later, for example segmentation-ready, aggregation-rich, or poor ice quality.",
            ),
            image_count=2,
            captions=(
                "Figure 11. Tomogram Gallery with filters, thumbnail grid, and details sidebar.",
                "Figure 12. Example of cumulative tag filtering and multi-selection in the gallery.",
            ),
            tips=(
                "Ratings are useful for broad quality ranking, while tags are better for categorical observations such as broken lamella, aggregates, good membranes, or contamination.",
                "Try to settle on a small controlled vocabulary for tags early in a project. This makes later filtering much more powerful.",
            ),
        ),
        WalkthroughSection(
            title="5. Running TS-based downstream jobs",
            goal="Run per-tilt-series jobs on a curated subset instead of on the whole dataset.",
            steps=(
                "Open Processing: TS jobs and confirm that the TS processing list contains the tomograms you want to work on.",
                "If the list is incomplete, add more tilt series either directly in this tab or by returning to the Tomogram Gallery.",
                "Choose the TS job type you want to run, for example a segmentation, denoising, or template-matching task.",
                "Review the command preview and confirm that CryoPal_tomo resolved the expected TS-specific inputs automatically.",
                "Choose whether to run locally, submit to Slurm, or schedule the job for later execution.",
                "If you want to operate from history, switch the job selector to the history view and launch the scheduled entries from there.",
            ),
            checks=(
                "This tab is most efficient when the TS processing list has already been curated in the gallery.",
            ),
            image_count=2,
            captions=(
                "Figure 13. TS processing list populated from curated tomograms.",
                "Figure 14. Example TS job configuration and command preview.",
            ),
            tips=(
                "Think of the TS processing list as your active worklist. If too many tilt series are in it, clear it and rebuild a smaller, more intentional selection.",
                "The main advantage of this list is that you define the subset once and can then reuse it across several TS jobs without repeatedly browsing for the same tomograms.",
            ),
        ),
        WalkthroughSection(
            title="6. Extending CryoPal_tomo with the File Registry and Custom jobs",
            goal="Teach CryoPal_tomo about new file types and then reuse them inside custom workflows.",
            steps=(
                "Open File registry and inspect how CryoPal_tomo currently resolves built-in roles such as tomograms, angle files, aligned stacks, and MDOCs.",
                "Click Add file role and define a new role for a file type CryoPal_tomo does not yet know, for example membrane segmentations or mitochondria masks.",
                "Fill in the directory template, filename pattern, recursive behavior, and TS-matching settings until the bottom association list shows the expected files for the expected tilt series.",
                "Return to Processing: Custom jobs and choose Build custom job type if you want to use those newly resolved files as inputs to an in-house script.",
                "Define the custom command, create parameter rows, and choose file-registry-backed TS-selection input types where appropriate.",
                "Save the job and reopen it from the runtime view to confirm that the TS processing list now drives the resolved inputs automatically.",
                "Go back to Tomogram Gallery and double-click a thumbnail if you want to confirm that the new file role now appears in the associated-file list as well.",
            ),
            checks=(
                "Only add a new custom role once you can reliably predict its directory and filename pattern; this keeps later automation stable.",
                "Use Settings > Configure viewer defaults if the new file role should open with a specific external program rather than the system default.",
            ),
            image_count=10,
            captions=(
                "Figure 15. File Registry overview with existing file roles.",
                "Figure 16. Adding a new file role.",
                "Figure 17. Pattern definition for a segmentation-style role.",
                "Figure 18. TS associations resolved for the new role.",
                "Figure 19. Custom-job builder with file-registry-backed parameters.",
                "Figure 20. Management dialog for custom job types.",
                "Figure 21. Saved custom job shown in the runtime tab.",
                "Figure 22. Gallery details with newly associated custom file roles.",
                "Figure 23. Viewer-default configuration for preferred external tools.",
                "Figure 24. Additional viewer/defaults example for role-specific opening logic.",
            ),
            tips=(
                "Only add a new custom file role once you can state clearly where the files live and how their names relate to TS identifiers. Ambiguous file roles create confusion everywhere they are reused.",
                "Custom jobs are most useful when their inputs can be made dataset-aware or TS-aware. If a script still requires you to browse manually for every run, it has not yet benefited fully from CryoPal_tomo's structure.",
            ),
        ),
        WalkthroughSection(
            title="7. Keeping paths, exports, and shared settings under control",
            goal="Use CryoPal_tomo's maintenance tools to validate project state and to export reusable records.",
            steps=(
                "Open Settings > Check paths whenever you want a fast overview of what CryoPal_tomo can and cannot currently resolve.",
                "Use Show details to inspect missing files at the dataset level or at the individual tilt-series level.",
                "Open File > Export job history if you need a CSV-based record of the executed or submitted processing actions.",
                "Open File > Export file paths if another software tool needs a clean list of files that CryoPal_tomo already knows about.",
                "Use Settings > Export .cryopal.settings-file when you want to share selected defaults, environments, custom jobs, or viewer settings with another user or another project.",
                "Use the corresponding import dialog when you want to adopt only certain settings categories without replacing the whole project file.",
            ),
            checks=(
                "A project file stores the working state of one project. A .cryopal.settings file is better suited for sharing reusable configuration patterns across projects.",
            ),
            image_count=7,
            captions=(
                "Figure 25. Check paths summary dialog.",
                "Figure 26. Detailed path report for datasets and file roles.",
                "Figure 27. Export job history dialog.",
                "Figure 28. Export file paths dialog.",
                "Figure 29. Settings-bundle export selection.",
                "Figure 30. Settings-bundle import selection.",
                "Figure 31. Overwrite handling during settings import.",
            ),
            tips=(
                "Check paths is especially helpful after moving data, renaming directories, importing legacy datasets, or changing file-registry rules.",
                "Settings bundles are ideal for sharing lab conventions. Project files are better reserved for the state of a specific scientific project.",
            ),
        ),
        WalkthroughSection(
            title="8. Particle export and STAR-file utilities",
            goal="Use the particle tab for STAR-file export, utilities, and quick downstream analysis plots.",
            steps=(
                "Open Processing: Particle jobs and select Export particles if you want CryoPal_tomo to launch particle export across several datasets in one place.",
                "Add the datasets you want to include, then fill in the required export parameters and inspect the command preview before execution.",
                "Run the export locally or through Slurm, depending on your setup.",
                "Switch to the utility-style particle jobs when you need to distance-clean, intersect, merge, or split STAR files.",
                "Use Plot particle abundance when you want a quick condition-level or dataset-level comparison of particle counts or densities.",
                "Use Plot classification convergence when you want CryoPal_tomo to inspect a Relion classification folder and summarize iteration-to-iteration class behavior.",
            ),
            checks=(
                "If Save particle plots is enabled in Settings > Set preferences, rendered plots will remain available from job-history details.",
            ),
            image_count=4,
            captions=(
                "Figure 32. Particle export across multiple datasets.",
                "Figure 33. Particle-abundance job setup.",
                "Figure 34. Example particle-abundance plots.",
                "Figure 35. Example classification-convergence plots.",
            ),
            tips=(
                "Use particle plots as quick decision aids rather than as a replacement for deeper statistical analysis. They are excellent for spotting trends early.",
                "When comparing several STAR files, make sure you remain aware of whether you are comparing conditions, datasets, or differently filtered versions of the same underlying particle pool.",
            ),
        ),
        WalkthroughSection(
            title="9. Using shortcuts for recurring actions",
            goal="Speed up repetitive lab routines that do not need formal job tracking.",
            steps=(
                "Open the Shortcuts tab if you often launch the same software, move into the same working directory, or run the same short command sequence.",
                "Double-click the plus tile to create a new shortcut.",
                "Give the shortcut a clear title, enter the command sequence line by line, and choose a tile color.",
                "Double-click the finished tile to test the shortcut in its own log window.",
                "If you want to edit, clone, remove, import, or export shortcuts later, manage them through Settings > Manage shortcuts.",
            ),
            checks=(
                "Shortcuts are best for convenience actions. Use processing tabs when you need formal project history and reproducible job metadata.",
            ),
            image_count=3,
            captions=(
                "Figure 36. Shortcuts dashboard.",
                "Figure 37. Shortcut creation dialog.",
                "Figure 38. Shortcut-management dialog in Settings.",
            ),
            tips=(
                "Shortcuts are a good place for GUI launchers, environment activation plus one tool, or directory-jump helpers. Avoid turning them into full scientific pipelines if you still want clear project history.",
            ),
        ),
        WalkthroughSection(
            title="10. Moving into MTools and MCore",
            goal="Continue from STA outputs into MTools or MCore without leaving the CryoPal_tomo project context.",
            steps=(
                "Open Processing: M when you are ready to work with M populations rather than only with datasets or tilt series.",
                "Create a new M population or import an existing .population file.",
                "Confirm that the population summary, species, and source information are displayed correctly before launching jobs.",
                "Choose the MTools or MCore command you need, adjust the parameters, and then run, schedule, or submit the job just as you would in the other processing tabs.",
            ),
            checks=(
                "If local M execution depends on a specific environment in your setup, launch CryoPal_tomo from the correct environment or confirm that the configured environment selection is appropriate.",
            ),
            image_count=2,
            captions=(
                "Figure 39. Processing: M tab with population selection.",
                "Figure 40. Example M-job setup and history view.",
            ),
            tips=(
                "The M tab becomes much easier to use once populations, species, and sources are named consistently. Good naming reduces confusion later when several refinement branches exist.",
            ),
        ),
    )


def add_walkthrough(document: Document) -> None:
    add_part_heading(
        document,
        "Part III. Hands-on Walkthrough",
        "A practical, screenshot-based tour through a typical CryoPal_tomo-supported cryo-ET processing journey.",
    )

    images = natural_image_list()
    cursor = 0
    for section in walkthrough_sections():
        document.add_heading(section.title, level=1)
        document.add_paragraph(section.goal)
        for step in section.steps:
            document.add_paragraph(step, style="List Number")
        if section.checks:
            document.add_heading("What to check", level=2)
            for check in section.checks:
                document.add_paragraph(check, style="List Bullet")
        if section.tips:
            document.add_heading("Tips and tricks", level=2)
            for tip in section.tips:
                document.add_paragraph(tip, style="List Bullet")
        if section.image_count:
            for caption in section.captions:
                image_path = images[cursor]
                add_figure(document, image_path, caption)
                cursor += 1

    document.add_heading("Closing note", level=1)
    document.add_paragraph(
        "CryoPal_tomo continues to evolve. As with many research software projects, users may occasionally encounter "
        "rough edges or workflow-specific limitations. In practice, however, the software already provides a strong "
        "framework for keeping cryo-ET processing organized, reproducible, and easier to review across datasets and "
        "time. When bugs or missing features are discovered, documenting them and feeding them back into development "
        "helps strengthen the software for the broader user community."
    )


def build_document() -> Path:
    document = Document()
    configure_page(document)
    style_document(document)
    add_cover(document)
    add_prefatory_section(document)
    add_table_of_contents(document)

    images = natural_image_list()
    add_quick_start(document, images)
    add_conceptual_guide(document)
    add_walkthrough(document)
    add_header_footer(document)
    document.save(str(OUTPUT_DOCX))
    return OUTPUT_DOCX


if __name__ == "__main__":
    output = build_document()
    print(output)
